#!/usr/bin/env python3
"""
Genera la versione .docx del manuale DrumAPPBass v1.0 a partire dalla
struttura del manuale Markdown. Layout pulito con:
- intestazioni numerate H1-H4
- TOC automatico in cima (Word lo aggiorna a F9 / Update Field)
- page break tra sezioni principali
- palette coerente con l'app: avorio #eae3d2 (sfondo blockquote),
  inchiostro #1a1a22 (testo), arancio #f77f00 (heading H1)
- font: Calibri body, Calibri Light heading
- tabelle con header colorato

Output: manuali/MANUALE_DrumAPPBass_v1.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

# Palette
INK   = RGBColor(0x1A, 0x1A, 0x22)   # #1a1a22
PAPER = RGBColor(0xEA, 0xE3, 0xD2)   # #eae3d2
ORANGE= RGBColor(0xF7, 0x7F, 0x00)   # #f77f00
BASS  = RGBColor(0x2A, 0x3A, 0x5C)   # #2a3a5c
GREY  = RGBColor(0x55, 0x55, 0x5A)

OUT = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   'MANUALE_DrumAPPBass_v1.docx')

doc = Document()

# Margini
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Stili default
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = INK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

def style_heading(level, color, size, bold=True, space_before=12, space_after=6):
    s = styles[f'Heading {level}']
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = color
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)
    s.paragraph_format.keep_with_next = True

style_heading(1, ORANGE, 22)
style_heading(2, INK, 16)
style_heading(3, BASS, 13)
style_heading(4, INK, 12)

# Helper: page break
def page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

# Helper: heading with text
def H1(text):
    h = doc.add_heading(text, level=1)
    return h

def H2(text):
    h = doc.add_heading(text, level=2)
    return h

def H3(text):
    h = doc.add_heading(text, level=3)
    return h

def H4(text):
    h = doc.add_heading(text, level=4)
    return h

# Helper: paragraph with optional inline formatting via list of (text, opts)
def P(text, italic=False, bold=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if italic: r.italic = True
    if bold: r.bold = True
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

# Helper: paragraph with mixed runs (list of dicts: {text, bold, italic, mono, color})
def Pmix(parts):
    p = doc.add_paragraph()
    for part in parts:
        r = p.add_run(part['text'])
        if part.get('bold'): r.bold = True
        if part.get('italic'): r.italic = True
        if part.get('mono'):
            r.font.name = 'Consolas'
            r.font.size = Pt(10)
        if part.get('color'): r.font.color.rgb = part['color']
    return p

# Helper: blockquote / nota importante (sfondo carta avorio)
def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EAE3D2')  # paper
    pPr.append(shd)
    # bordo sinistro arancio
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), 'F77F00')
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(text)
    r.italic = True
    return p

# Helper: bullet
def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

# Helper: codice
def code(text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F2EEDF')
    pPr.append(shd)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9.5)
    r.font.color.rgb = INK
    return p

# Helper: tabella (header riga + body rows)
def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light List Accent 1'  # base, override colors below
    t.autofit = True
    # header
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = PAPER
        r.font.size = Pt(10.5)
        # sfondo cella header inchiostro
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), '1A1A22')
        tcPr.append(shd)
    # body
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cell = cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.color.rgb = INK
            # zebra
            if ri % 2 == 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F7F3E5')
                tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t

# Helper: TOC field code (Word lo aggiorna premendo F9)
def insert_toc():
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = 'Aggiorna l\'indice con F9 (Word) o Cmd+Opt+Shift+U (Mac)'
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    run._r.append(fldChar4)

# Helper: paragrafo con testo monospaziato inline (tasti tipo `B`)
def Pkbd(text):
    """Sostituisce `xxx` con run monospaziato leggermente colorato (chiave)."""
    p = doc.add_paragraph()
    parts = []
    cur = ''
    in_kbd = False
    for ch in text:
        if ch == '`':
            if cur:
                parts.append((cur, in_kbd))
            cur = ''
            in_kbd = not in_kbd
        else:
            cur += ch
    if cur:
        parts.append((cur, in_kbd))
    for txt, kbd in parts:
        r = p.add_run(txt)
        if kbd:
            r.font.name = 'Consolas'
            r.font.size = Pt(10.5)
            r.font.color.rgb = BASS
            r.bold = True
    return p

# ============================================================
# COVER
# ============================================================

# Titolo cover
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(36)
r = title.add_run('Manuale utente')
r.font.size = Pt(28); r.font.color.rgb = INK; r.bold = True
title.add_run('\n')
r = title.add_run('DrumAPPBass v1.0')
r.font.size = Pt(36); r.font.color.rgb = ORANGE; r.bold = True

# Sub
sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(24)
r = sub.add_run('Edizione di anteprima riservata')
r.font.size = Pt(14); r.italic = True; r.font.color.rgb = BASS
sub.add_run('\n')
r = sub.add_run('Weekend App — not Public')
r.font.size = Pt(12); r.italic = True; r.font.color.rgb = GREY

# Metadati cover
doc.add_paragraph()
doc.add_paragraph()
meta_table = doc.add_table(rows=8, cols=2)
meta_table.autofit = True
meta_rows = [
    ('Versione manuale', '1.0'),
    ('Versione applicazione', '1.0 (Round 7 — Performance Modes)'),
    ('Data', '25 aprile 2026'),
    ('Autore', 'Alessandro Pezzali — pezzaliapp'),
    ('Sito', 'pezzaliapp.com'),
    ('Repository', 'github.com/pezzaliapp/DrumAPPBass'),
    ('App live', 'pezzaliapp.github.io/DrumAPPBass'),
    ('Licenza', 'MIT'),
]
for i, (k, v) in enumerate(meta_rows):
    cells = meta_table.rows[i].cells
    cells[0].text = ''
    cells[1].text = ''
    rk = cells[0].paragraphs[0].add_run(k)
    rk.bold = True; rk.font.color.rgb = INK; rk.font.size = Pt(10.5)
    rv = cells[1].paragraphs[0].add_run(v)
    rv.font.color.rgb = INK; rv.font.size = Pt(10.5)
    cells[0].width = Cm(5.5)
    cells[1].width = Cm(11)

# Disclaimer cover
doc.add_paragraph()
doc.add_paragraph()
quote(
    'Documento distribuito in via esclusiva alla community early-adopter dell\'autore '
    'prima del rilascio pubblico. Contenuto fattualmente verificato sul codice del repository '
    'al commit 107cff3 (25 aprile 2026).'
)

page_break()

# ============================================================
# INDICE
# ============================================================
H2('Indice')
insert_toc()

page_break()

# ============================================================
# 1. INTRODUZIONE
# ============================================================
H1('1. Introduzione')

H2('1.1 Cos\'è DrumAPPBass')
P(
'DrumAPPBass è una drum machine + bass synth che gira interamente nel browser. È un fork '
'esteso di DrumAPP e aggiunge una sezione bass completamente sintetizzata, uno step '
'sequencer dedicato con accent e slide, un live looper con tastiera on-screen, quattro '
'Performance Modes che cambiano il carattere globale del suono, e una libreria di '
'trentadue demo importabili.'
)
P(
'L\'applicazione è una PWA: gira offline dopo la prima visita, si installa come app sul '
'telefono o sul desktop, e non richiede registrazione né account. Tutto il suono è generato '
'in tempo reale via Web Audio API. Non ci sono sample WAV, non ci sono librerie esterne, '
'non ci sono dipendenze npm. È un singolo file app.js di circa 4000 righe, un index.html, '
'un foglio di stile e un service worker.'
)

H2('1.2 Filosofia')
P('Tre principi attraversano tutto il progetto.')

H4('Zero asset, vanilla Web Audio')
P(
'Ogni voce drum e bass è costruita con OscillatorNode, GainNode, BiquadFilterNode, '
'WaveShaperNode. Non si scarica nulla oltre al codice dell\'app. La conseguenza è '
'duplice: il pacchetto resta sotto i 200 KB e ogni utente ascolta esattamente lo stesso '
'suono in modo deterministico, indipendentemente dal browser o dal dispositivo.'
)

H4('Sketch musicale, non DAW professionale')
P(
'DrumAPPBass è progettato per la fase in cui un\'idea ritmica o una bassline va appuntata '
'velocemente, su una panchina o in un treno. Non sostituisce Ableton o Logic; integra una '
'fase precedente — quella in cui un\'idea, se non viene fissata in trenta secondi, va persa.'
)

H4('Non rompere niente')
P(
'Ogni feature aggiunta in un Round successivo non deve compromettere il funzionamento di '
'quelle precedenti. Le demo storiche di DrumAPP continuano a caricarsi e a suonare '
'identiche; le cinque demo bass curate del v1.0 base continuano a suonare in modalità '
'CLASSIC anche dopo l\'introduzione delle Performance Modes.'
)

H2('1.3 Posizionamento Weekend App')
P(
'DrumAPPBass fa parte della serie Weekend App — not Public, una linea di applicazioni web '
'sviluppate dall\'autore nei fine settimana e distribuite in due fasi.'
)
P(
'La prima fase è una anteprima riservata alla community che segue l\'autore su WhatsApp e '
'sui canali social. Gli early adopter ricevono il link dell\'applicazione insieme a questo '
'manuale, possono usare l\'app prima del rilascio pubblico e hanno un canale diretto per dare '
'feedback e segnalare bug. Questo documento accompagna proprio questa fase.'
)
P(
'La seconda fase è il rilascio pubblico sul sito dell\'autore (pezzaliapp.com) '
'e sulle directory pubbliche di applicazioni web. La data del rilascio '
'pubblico non è dichiarata in anticipo; dipende dal feedback raccolto in fase di anteprima '
'e dal completamento di eventuali raffinamenti.'
)

H2('1.4 A chi si rivolge')
P(
'Il manuale assume che il lettore abbia familiarità con concetti di base della produzione '
'musicale (BPM, swing, velocity, kick, snare). Non assume conoscenze avanzate di sintesi '
'sonora né esperienza con DAW professionali. Le sezioni tutorial sono pensate per chi '
'parte da zero; l\'appendice tecnica è per chi vuole capire l\'architettura interna o '
'estendere il codice.'
)

page_break()

# ============================================================
# 2. QUICK START
# ============================================================
H1('2. Quick Start — cinque minuti')

H2('2.1 Aprire l\'app')
P(
'Apri l\'URL dell\'applicazione nel browser. Su mobile, dopo qualche secondo il browser '
'proporrà di aggiungere l\'app alla schermata home: confermando, l\'app si installerà come '
'PWA e da quel momento si aprirà in modalità standalone (senza barra del browser) e '
'funzionerà offline.'
)
P(
'Al primo tocco o click qualsiasi, il browser sblocca il contesto audio. Da qui in poi '
'l\'applicazione è pronta a suonare.'
)

H2('2.2 Caricare una demo')
Pkbd(
'Premi il bottone DEMOS nel footer (oppure il tasto rapido `D` per caricare la demo '
'predefinita). Si apre una modale con la libreria delle trentadue demo divise per '
'categoria. Scegli per esempio Synth-Pop (WIDE) dalle demo di tipo perf (icona ⚡, '
'sfondo arancio chiaro): la demo carica pattern, parametri, master mixer, song '
'sequence e modalità performance del basso.'
)

H2('2.3 Suonare')
Pkbd(
'Premi `SPACE` o il bottone PLAY per avviare la riproduzione. La griglia si illumina '
'colonna per colonna seguendo lo step corrente; il drum suona in alto, il bass in basso. '
'Premendo `1`, `2`, `3`, `4` si cambia pattern al volo (A, B, C, D); con le parentesi '
'quadre `[` e `]` si scorre indietro e avanti tra i pattern.'
)

H2('2.4 Modificare')
Pkbd(
'Tocca le celle della griglia drum per accendere o spegnere uno step. La traccia attiva è '
'quella evidenziata nel pannello EDITING in alto: per cambiarla, click sul nome (KICK, '
'SNARE, eccetera) o usa le frecce `↑` `↓`.'
)
Pkbd(
'Per il basso premi il tasto `B`: il pannello in alto mostra ora i parametri del basso '
'(volume, pan, cutoff, resonance, envelope, decay, drive). La griglia bass è subito sotto '
'quella drum: tocca una cella per accendere uno step (la prima volta entra come `C2`), '
'drag verticale per cambiare la nota.'
)

H2('2.5 Salvare e condividere')
P(
'Per salvare lo stato corrente in uno slot locale, tieni premuto uno dei quattro '
'bottoni A/B/C/D nella sezione SLOT per mezzo secondo. Il bottone si illumina e flash '
'conferma. Per ricaricare lo stesso stato in seguito, basta un tocco breve sul bottone.'
)
P(
'Per esportare il pattern come file JSON premi EXPORT; per condividere via link SHARE '
'(l\'URL viene copiato negli appunti).'
)

page_break()

# ============================================================
# 3. INTERFACCIA
# ============================================================
H1('3. Interfaccia generale')

H2('3.1 Layout')
P('La schermata è divisa in fasce orizzontali, dall\'alto verso il basso.')
table(
    ['Fascia', 'Contenuto'],
    [
        ('Header', 'Brand, transport (PLAY/STOP), tempo (BPM + TAP), swing, master mixer (DRUM e BASS)'),
        ('Active track panel', 'Parametri della traccia attiva (drum o bass a seconda del focus)'),
        ('Sequencer drum', 'Griglia 8 voci × 16 step (o lunghezza scelta)'),
        ('Bass panel', 'Selettore STEP/LOOPER, EDIT, PERF mode, griglia bass, tastiera on-screen'),
        ('Modebar', 'Edit mode drum, pattern A/B/C/D, song toggle, lunghezza pattern, humanize, metronomo'),
        ('Step readout', 'Hint contestuale dello step e della cella attiva'),
        ('Songbar', 'Sequence editor visuale (slot di pattern A/A/B/A…)'),
        ('Storage', 'Slot locali, file IO, output (BOUNCE/REC/MIDI/SHARE), history (COPY/PASTE/UNDO/REDO)'),
        ('Footer', 'DEMO, DEMOS, CLEAR, HELP e shortcut summary'),
    ],
    col_widths=[4.5, 12]
)

H2('3.2 Palette e codici colore')
P(
'L\'estetica è quella della serie Studio Press dell\'autore: carta avorio #eae3d2, '
'inchiostro #1a1a22, accento arancione #f77f00 per i call-to-action principali. La sezione '
'bass ha un accento complementare blu notte #2a3a5c per distinguerla visivamente dalla '
'drum machine. Il selettore Performance Mode usa l\'arancio del brand: il bottone attivo '
'si riempie di arancio per ricordare visivamente che si tratta di un selettore di carattere '
'globale.'
)

H2('3.3 Responsive')
P(
'Su mobile il layout collassa: i pannelli si impilano verticalmente, i bottoni delle '
'modebar si rimpiccoliscono ma restano cliccabili. La tastiera on-screen del basso resta '
'utilizzabile in entrambi gli orientamenti. Per uso intensivo è consigliato il landscape.'
)

page_break()

# ============================================================
# 4. DRUM
# ============================================================
H1('4. Drum Machine')

H2('4.1 Le otto voci')
P('Le otto tracce drum sono sempre nello stesso ordine, dall\'alto verso il basso.')
table(
    ['#', 'Nome', 'Sintesi'],
    [
        ('1', 'KICK',    'Sine wave 165 → 45 Hz con envelope esponenziale + click 1200 Hz a transient'),
        ('2', 'SNARE',   'Triangle 220 Hz + noise high-pass 1 kHz, envelope indipendenti'),
        ('3', 'HI-HAT',  'Noise high-pass 7 kHz (chiuso, decay corto)'),
        ('4', 'OPEN HH', 'Noise high-pass 7 kHz, decay lungo'),
        ('5', 'CLAP',    'Noise multi-burst con envelope a quattro picchi'),
        ('6', 'TOM',     'Sine 110 Hz con sweep'),
        ('7', 'RIMSHOT', 'Noise band-pass + click'),
        ('8', 'COWBELL', 'Due square wave intervallate (560 Hz + 845 Hz)'),
    ],
    col_widths=[1, 2.5, 13]
)
P(
'Tutte le voci sono generate in tempo reale e rispondono ai parametri di traccia '
'(volume, pan, pitch, decay, filter, mute, solo).'
)

H2('4.2 Il sequencer')
P(
'La griglia ha sedici step di default. La lunghezza del pattern è regolabile fra otto, '
'dodici, sedici, ventiquattro e trentadue step (selettore LEN nella modebar). Il numero '
'di step visibile cambia immediatamente al variare di patternLength.'
)
P(
'Ogni cella della griglia rappresenta uno step della traccia corrispondente. Tocca o '
'clicca per accendere lo step; tocca di nuovo per spegnerlo. La cella accesa mostra una '
'piccola barra colorata che riflette il valore di velocity.'
)

H2('4.3 Edit mode')
P('Sopra il sequencer una modebar dedicata permette di cambiare il significato di drag e tap sulle celle.')
table(
    ['Mode', 'Cosa fa'],
    [
        ('TRIG',  'Tap = accende/spegne. Default.'),
        ('VEL',   'Drag verticale su cella accesa → velocity 0,05 – 1,0'),
        ('PROB',  'Drag verticale → probabilità in percentuale che lo step suoni quel ciclo'),
        ('RATCH', 'Drag verticale → ripetizioni interne 1, 2, 3, 4 (ratchet)'),
        ('NUDGE', 'Drag verticale → micro-timing ±50 ms fuori griglia'),
    ],
    col_widths=[3, 13.5]
)
P(
'Velocity, probability, ratchet e nudge convivono sulla stessa cella: una cella può essere '
'vel=0.85, prob=80, ratch=2, nudge=+12. Tutti e quattro i valori sono persistiti nei file '
'di salvataggio.'
)

H2('4.4 Pattern multipli e song mode')
Pkbd(
'Quattro pattern indipendenti A, B, C, D coesistono nello stesso set. Si commutano al volo '
'durante il playback con i bottoni nella modebar PATTERN, con i tasti `1`–`4` o con `[` e '
'`]`. Il pattern attivo è quello evidenziato.'
)
P(
'L\'editor SEQUENCE in basso permette di costruire una song concatenando pattern: per '
'esempio A A B A C C. Click sul singolo slot della sequence cycle attraverso A → B → C → '
'D → A. I bottoni − e + rimuovono o aggiungono slot (minimo 1, massimo 16). Quando il '
'toggle SONG è attivo, la riproduzione segue la sequence invece di restare sul pattern '
'singolo.'
)

H2('4.5 Parametri di traccia')
Pkbd(
'Premendo il nome di una traccia (KICK, SNARE…) la traccia diventa attiva e il pannello '
'EDITING in alto mostra i suoi parametri. Si può anche cambiare traccia con `↑` `↓`.'
)
table(
    ['Parametro', 'Range', 'Note'],
    [
        ('VOL',    '0 – 100',                 'Volume di traccia'),
        ('PAN',    '−100 – +100',              'Doppio click per centro'),
        ('PITCH',  '−12 – +12 semitoni',       'Trasposizione cromatica'),
        ('DECAY',  '0,4× – 2,5×',              'Moltiplicatore del decay envelope'),
        ('FILTER', 'OFF / LOW-PASS / HIGH-PASS', 'Filtro per voce'),
        ('CUTOFF', '0 – 100',                  'Frequenza di taglio (dipende dal tipo)'),
    ],
    col_widths=[3, 5, 8.5]
)
P(
'I parametri di traccia sono condivisi tra tutti e quattro i pattern: cambiano il timbro '
'della traccia in modo globale, non per pattern.'
)

H2('4.6 Mute e solo')
Pkbd(
'`M` mette in mute la traccia attiva, `S` la mette in solo. Quando una traccia è in solo, '
'le altre vengono temporaneamente silenziate; in modalità solo, anche il basso viene '
'silenziato a meno che il basso stesso sia in solo.'
)

H2('4.7 Master mixer')
P(
'Nell\'header in alto a destra ci sono due slider master: uno per il bus drum, uno per il '
'bus bass. Sono indipendenti e si applicano a valle di ogni traccia ma a monte del master '
'globale, quindi vengono rispettati dal bounce WAV e dal REC live. Il valore di default è '
'90 per il drum e 80 per il bass; doppio click ripristina il default.'
)

page_break()

# ============================================================
# 5. BASS SYNTH
# ============================================================
H1('5. Bass Synth')

H2('5.1 Architettura')
P(
'La sezione bass è una sintesi sottrattiva ibrida ispirata alla TB-303 ma più moderna nel '
'timbro. Per ogni nota suonata, l\'app costruisce in tempo reale un grafo di nodi Web Audio:'
)
bullet('uno stack di sawtooth detunate (cinque voci di default in modalità CLASSIC, da tre a cinque a seconda della Performance Mode) che genera la massa armonica principale;')
bullet('un sub-oscillatore square a un\'ottava sotto, attivo solo quando la nota è sopra una soglia minima (di default C2, abbassata a B1 in modalità SUB) per evitare battimenti con il fundamental del kick;')
bullet('un mixer saw/sub con rapporto variabile a seconda della modalità;')
bullet('un gate envelope di ampiezza con attack, sustain, release configurabili;')
bullet('un lowpass biquad modulato per nota da un envelope ADSR del filtro, con Q dinamica scalata su velocity, accent e nota suonata;')
bullet('un WaveShaper con curva tanh per il soft-clipping (drive);')
bullet('un panner globale e un gain di traccia che alimentano il bus bass.')
P(
'Il bus bass passa per un high-pass a 30 Hz prima del master, per rimuovere ogni residuo '
'sub-audio (il fundamental del basso E1 a 41 Hz resta intatto).'
)

H2('5.2 STEP mode — il sequencer del basso')
P(
'In modalità STEP il basso ha la sua griglia da sedici celle (sincronizzata con il drum). '
'Ogni cella accesa è una nota con cinque parametri.'
)
table(
    ['Parametro', 'Significato', 'Range'],
    [
        ('note',   'Nota suonata',                                    'C1 – B3'),
        ('vel',    'Velocity',                                        '0,05 – 1,0'),
        ('len',    'Length / gate',                                   '10 – 100 % dello step'),
        ('accent', 'Boolean',                                         '+6 dB e filtro più aperto'),
        ('slide',  'Boolean — portamento 30 ms verso la nota successiva', 'tie-note in stile 303'),
    ],
    col_widths=[3, 7, 6.5]
)
P(
'Il modello è quello del 303: una sequenza monofonica dove i due parametri timbrici '
'dominanti sono accent e slide. L\'accent alza il picco iniziale a 1,2× e raddoppia '
'l\'envelope del filtro su quel colpo. Lo slide non ri-triggera l\'envelope del filtro e '
'applica un glide lineare di trenta millisecondi sulla frequenza, creando il classico "tie" '
'della TB-303.'
)

H2('5.3 EDIT mode bass')
P('Sopra la griglia bass una modebar dedicata permette di scegliere quale parametro modificare con drag o tap.')
table(
    ['Mode', 'Comportamento'],
    [
        ('NOTE',  'Tap su cella spenta → accende a C2. Tap su accesa → spegne. Drag verticale → cambia la nota. Long-press di mezzo secondo apre un mini-selettore con tutti i semitoni C1–B3. Scroll della rotella su desktop.'),
        ('VEL',   'Drag verticale su cella accesa → velocity. Tap su spenta non fa nulla.'),
        ('LEN',   'Drag verticale → length / gate.'),
        ('ACC',   'Tap su cella accesa toggla accent.'),
        ('SLIDE', 'Tap su cella accesa toggla slide. Il bordo mostra una freccia → verso lo step successivo.'),
    ],
    col_widths=[2.5, 14]
)

H2('5.4 LOOPER mode')
P('Il toggle STEP / LOOPER nel bass panel cambia il modo di registrazione del basso. In LOOPER:')
bullet('la griglia step diventa read-only (in grigio) e mostra eventualmente le palline delle note registrate live;')
bullet('si attiva la tastiera on-screen (un\'ottava con bottoni OCT − e OCT +);')
Pkbd('si attivano le shortcut pianistiche del computer: `A W S E D F T G Y H U J K` mappano i dodici semitoni di un\'ottava (più la C alta su `K`); `Z` e `X` cambiano l\'ottava;')
Pkbd('il bottone REC (o tasto `R`) arma la registrazione: il rec parte sul prossimo downbeat e cattura nota + timestamp frazionario fino a fine pattern, poi si rinnova in overdub mode;')
bullet('l\'opzione QUANT 1/16 quantizza i timestamp al sedicesimo più vicino;')
bullet('l\'opzione PLAY STEP TOO fonde step pattern e live loop (di default in LOOPER si suonano solo le note del loop).')
Pkbd(
'I tasti pianistici sulla tastiera del computer sono attivi solo in modalità LOOPER per '
'non collidere con le shortcut globali (`T` tap tempo, `S` solo, `D` demo, eccetera).'
)
P(
'Ogni pattern A/B/C/D ha due slot indipendenti: uno step pattern e un live loop. Si '
'possono usare uno, l\'altro, o entrambi.'
)

H2('5.5 Parametri di traccia bass')
Pkbd('Premendo `B` il pannello in alto mostra i parametri del basso al posto di quelli della drum.')
table(
    ['Parametro', 'Range', 'Descrizione'],
    [
        ('VOL',    '0 – 100',     'Volume di traccia bass'),
        ('PAN',    '−100 – +100', 'Posizione stereo'),
        ('CUTOFF', '0 – 100',     'Frequenza di taglio del filtro (50 Hz – 5 kHz, esponenziale)'),
        ('RES',    '5 – 150',     'Q del filtro (0,5 – 15)'),
        ('ENV',    '−100 – +100', 'Modulazione del filtro tramite envelope'),
        ('DECAY',  '60 – 800 ms', 'Tempo di decay dell\'envelope del filtro'),
        ('DRIVE',  '0 – 100',     'Quantità di soft-clipping WaveShaper'),
    ],
    col_widths=[2.5, 4, 10]
)
P(
'Tutti i parametri possono essere automatizzati durante il REC live: ogni movimento dei '
'knob viene catturato nel webm/opus.'
)

H2('5.6 Routing MIDI')
P(
'Le note bass vengono inviate sul canale MIDI 2; le drum restano sul canale 10 GM standard '
'(kick=36, snare=38, hihat=42, openhat=46, clap=39, tom=45, rim=37, cow=56). Si può quindi '
'triggerare uno strumento esterno via MIDI con il sequencer del basso senza interferenze '
'con le drum.'
)

page_break()

# ============================================================
# 6. PERFORMANCE MODES
# ============================================================
H1('6. Performance Modes')

H2('6.1 Idea')
Pkbd(
'Il tasto `P` (e i quattro bottoni arancio nel bass panel) cambia il carattere globale del '
'basso. Le quattro modalità sono scenari coerenti: ognuna ridefinisce insieme '
'l\'architettura della sintesi (numero di voci, spread stereo, detune, soglia del '
'sub-oscillatore) e i parametri di timbro (drive, cutoff, resonance, forma dell\'envelope).'
)
P(
'L\'utente non sente variazioni sottili: ogni modalità produce un suono riconoscibile e '
'diverso. Questo è un selettore di scenario, non un tweak.'
)

H2('6.2 Le quattro modalità')
table(
    ['Modalità', 'Architettura', 'Timbro'],
    [
        ('CLASSIC', 'mono · 5 voci sawtooth ±10c · sub a –12 st al 30 % · soglia C2',
                    'baseline Round 6, suono bilanciato per tutti i generi'),
        ('WIDE',    'stereo · 5 voci ±18c pannate –1 / –0,5 / 0 / +0,5 / +1 · sub al 25 %',
                    'drive ×0,85, sustain alta, immersivo'),
        ('PUNCH',   'mono · 5 voci ±5c · sub al 45 % · soglia C2',
                    'drive ×1,5, decay rapido, sustain 0,55, release 25 ms'),
        ('SUB',     'mono · 3 voci ±5c · sub al 65 % (primario) · soglia abbassata a B1',
                    'drive ×1,6, sustain 0,85, sub-heavy'),
    ],
    col_widths=[2.5, 7, 7]
)

H2('6.3 Quando usare quale')
P(
'CLASSIC è la scelta neutra. È la baseline del Round 6, e tutte le demo bass storiche '
'(funk, house, onedrop, boombap, trap) suonano in CLASSIC. Da usare quando il basso deve '
'stare in pocket senza colorare lo scenario.'
)
P(
'WIDE brilla quando il basso è in registro medio (sopra C2) e ha un ruolo melodico. Si '
'sente immediatamente in cuffia: lo stereo spread crea un wash dove il basso "respira" da '
'sinistra a destra. Generi tipici: synth-pop, electro, synthwave, italo-disco, new wave.'
)
P(
'PUNCH è la scelta per i groove serrati. Il drive alto e la sustain corta producono un '
'basso percussivo che dialoga col kick senza sovrapporsi. Generi tipici: funk, neo-soul, '
'hip-hop dirty, boom bap, breakbeat.'
)
P(
'SUB sostituisce le voci sawtooth con il sub-oscillatore primario al 65 %. È il preset per '
'i generi sub-heavy: dub, dubstep, DnB rolling, trap moderna con 808 lunghi. La soglia del '
'sub è abbassata a B1 per permettere note più gravi senza perdere il basso fondamentale.'
)

H2('6.4 Switching live')
Pkbd(
'Il cambio di modalità si può fare in qualsiasi momento, anche durante la riproduzione. '
'Premendo `P` si cycla CLASSIC → WIDE → PUNCH → SUB → CLASSIC; un toast a video conferma '
'la modalità attiva. La voce eventualmente in corso viene chiusa con un fade di dieci '
'millisecondi per evitare crepe sul cambio di detune o pan.'
)

H2('6.5 Persistenza')
P(
'La modalità performance è salvata insieme al resto dello stato negli slot locali, nel '
'JSON v3 (campo performanceMode al root) e nello share link (in via implicita: il pattern '
'si carica e poi la modalità viene ripristinata dal JSON).'
)

page_break()

# ============================================================
# 7. DEMO LIBRARY
# ============================================================
H1('7. Demo Library')

H2('7.1 Catalogo')
P('Il bottone DEMOS apre la libreria delle demo importabili: trentadue set drum (e drum + bass) divisi per categoria.')
table(
    ['Categoria', 'Numero', 'Tipo', 'Tag visivo'],
    [
        ('Modern',            '7', 'Drum-only',                  '🎛 grigio'),
        ('Iconic',            '9', 'Drum-only',                  '⭐ giallo'),
        ('Classic break',     '7', 'Drum-only',                  '🏛 marrone'),
        ('Bass curate',       '5', 'Drum + bass (CLASSIC)',      '🎸 blu notte'),
        ('Performance modes', '4', 'Drum + bass (mode-aware)',   '⚡ arancio'),
    ],
    col_widths=[5, 2, 6, 3.5]
)

H2('7.2 Demo modern')
table(
    ['File', 'Nome', 'BPM', 'Note'],
    [
        ('demo-house.json',         'House / Techno',     '124', '4/4 dritto, build + drop'),
        ('demo-trap.json',          'Trap moderno',       '140', '808 + ratchet hats'),
        ('demo-boombap.json',       'Boom Bap anni 90',    '90', 'Swing 52% Dilla style'),
        ('demo-dnb.json',           'DNB / Amen break',   '170', 'Jungle, ghost snare al 7'),
        ('demo-makesomenoise.json', 'NYC Hip-Hop 2011',   '105', 'Kick doppio, cowbell'),
        ('demo-ukhardcore.json',    'UK Hardcore 92-93',  '140', 'Pre-jungle rave'),
        ('demo-onedrop.json',       'Dub / Reggae',        '80', 'Beat sul 3, filosofia invertita'),
    ],
    col_widths=[5, 4, 1.5, 6]
)

H2('7.3 Demo iconic')
table(
    ['File', 'Nome', 'BPM'],
    [
        ('demo-dadada.json',        'Da Da Da-style',     '120'),
        ('demo-wewillrockyou.json', 'We Will Rock You',    '81'),
        ('demo-sevennation.json',   'Seven Nation Army',  '124'),
        ('demo-anotherone.json',    'Another One-style',  '110'),
        ('demo-superstition.json',  'Superstition-style', '100'),
        ('demo-rosanna.json',       'Rosanna Shuffle',     '87'),
        ('demo-stayinalive.json',   "Stayin' Alive",      '103'),
        ('demo-takefive.json',      'Take Five (5/4!)',   '172'),
        ('demo-wipeout.json',       'Wipe Out',           '162'),
    ],
    col_widths=[5, 7, 2]
)

H2('7.4 Demo classic break')
table(
    ['File', 'Nome', 'BPM'],
    [
        ('demo-billiejean.json',       'Billie Jean-style',     '117'),
        ('demo-funkydrummer.json',     'Funky Drummer-style',   '103'),
        ('demo-levee.json',            'Levee Breaks-style',     '72'),
        ('demo-apache.json',           'Apache-style',          '112'),
        ('demo-impeach.json',          'Impeach-style',         '100'),
        ('demo-ashleysroachclip.json', "Ashley's Roachclip",    '100'),
        ('demo-synthsub.json',         'Synthetic Sub-style',    '91'),
    ],
    col_widths=[5.5, 6.5, 2]
)

H2('7.5 Demo bass curate')
P('Caricano in modalità CLASSIC. Sono le demo originali della release Round 6, scritte con dialogo kick/basso ritmicamente sfalsato.')
table(
    ['File', 'Genere', 'BPM', 'Tonalità'],
    [
        ('demo-bass-funk.json',    'Funk slap-style',      '108', 'E minor'),
        ('demo-bass-house.json',   'House 4/4',            '124', 'A minor'),
        ('demo-bass-onedrop.json', 'Reggae one-drop',       '80', 'A minor'),
        ('demo-bass-boombap.json', 'Hip-hop boom bap',      '90', 'D minor'),
        ('demo-bass-trap.json',    'Trap 808 sub',         '140', 'F minor'),
    ],
    col_widths=[5.5, 5, 1.5, 4]
)

H2('7.6 Demo performance modes')
P('Una demo per ogni Performance Mode introdotta nel Round 7. Ognuna dichiara la propria modalità nel JSON e applica automaticamente lo scenario.')
table(
    ['File', 'Genere', 'BPM', 'Tonalità', 'Mode'],
    [
        ('demo-bass-synthwave.json', 'Synthwave / italo-disco',   '110', 'B minor',  'CLASSIC'),
        ('demo-bass-synthpop.json',  '80s synth-pop / electro',   '120', 'C minor',  'WIDE'),
        ('demo-bass-neosoul.json',   'Neo-soul / hip-hop dirty',   '96', 'F# minor', 'PUNCH'),
        ('demo-bass-dub.json',       'Dub rolling / one-drop',     '92', 'G minor',  'SUB'),
    ],
    col_widths=[5.5, 5.5, 1.5, 2.2, 1.8]
)

H2('7.7 Caricare e scaricare le demo')
P(
'Click su una card della modale carica la demo immediatamente (sovrascrivendo il set '
'corrente). Per scaricare il file JSON di una demo, vai nella cartella examples/ del '
'repository GitHub: ogni demo è un file .json indipendente che può essere importato a '
'sua volta con il bottone IMPORT.'
)
quote(
'Tutte le demo sono ricostruzioni ritmico-armoniche generiche ispirate ai loro generi di '
'riferimento. Non sono trascrizioni né copie di brani originali; non riproducono melodie '
'protette da copyright. I nomi tipo "Billie Jean-style" o "Seven Nation Army" indicano lo '
'stile, non il brano.'
)

page_break()

# ============================================================
# 8. WORKFLOW
# ============================================================
H1('8. Workflow esempi')

H2('8.1 Costruire una traccia da zero')
def numbered(items):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph(t, style='List Number')

numbered([
'Apri l\'app, premi D per caricare la demo predefinita (un four-on-the-floor pulito) come punto di partenza, oppure premi CLEAR per partire vuoto.',
'Scegli il BPM. Per un techno usa 124–130, per un hip-hop 85–95, per un drum & bass 165–175.',
'Costruisci il pattern A: kick sui downbeat, snare sul 5 e 13, hi-hat sugli ottavi, eventuale clap.',
'Premi B per passare al basso. Aggiungi la tonica sul primo step, una quinta o una settima sul step 9, qualche ghost note e uno slide tra step 8 e 9.',
'Scegli una Performance Mode (P cycla): WIDE per i generi sintetici, PUNCH per i groove serrati, SUB per il sub-heavy.',
'Premi 2 per passare al pattern B, copia il pattern A (COPY poi PASTE su B) e modifica leggermente.',
'Costruisci la song nella sequence: per esempio A A A B A A B B. Attiva il toggle SONG.',
'Premi SPACE e ascolta. Se ti convince, salva con hold su uno slot A/B/C/D.',
'Esporta il .json o premi BOUNCE per ottenere un file .wav.',
])

H2('8.2 Ricreare un genere')
P('L\'approccio sano è caricare una demo del genere come scheletro, poi modificarla.')
bullet('Reggae one-drop: parti da Reggae One-drop (demo bass), tieni il kick sul 3, non aggiungere kick sul 1, lascia spazi vuoti nel basso, usa CUTOFF 0,40 e RES bassa.')
bullet('Trap 808: parti da Trap 808 Bass (demo bass), modalità SUB, basso in registro medio (F2 – C3), kick sui downbeat 1/7/11, ratchet 2× e 3× sull\'hi-hat, DRIVE 0,3.')
bullet('Funk: parti da Funk Em + Bass, modalità PUNCH, ghost notes ravvicinate sul basso, tonica solo sui downbeat principali, hi-hat con velocity alternata.')

H2('8.3 Live looper bass')
P('Per registrare al volo una bassline:')
numbered([
'Premi B per il focus bass, poi L per passare in LOOPER mode.',
'Premi SPACE per avviare il transport.',
'Premi R per armare la registrazione: REC partirà sul prossimo downbeat.',
'Suona la tastiera (mouse, touch o tasti A W S E D F T G Y H U J K); usa Z X per cambiare ottava.',
'Al termine del pattern, REC va in overdub: puoi aggiungere altre note al loop senza fermarti.',
'Premi di nuovo R per fermare la registrazione.',
'Per cancellare il loop corrente, CLEAR LOOP nel pannello.',
])
P('Il loop registrato è agganciato al pattern corrente: cambiando pattern, cambia anche il loop. Quattro pattern → quattro loop indipendenti.')

H2('8.4 Confrontare le Performance Modes')
P('Per percepire al meglio la differenza tra le quattro modalità:')
numbered([
'Carica Synth-Pop (WIDE) dalle demo perf.',
'Avvia la riproduzione.',
'Premi P ripetutamente per cyclare CLASSIC → WIDE → PUNCH → SUB.',
'Tieni le cuffie: WIDE produce uno spread stereo immediatamente percepibile, SUB cambia radicalmente il bilanciamento basso, PUNCH aggiunge mordente e accorcia la coda.',
])

page_break()

# ============================================================
# 9. ADVANCED
# ============================================================
H1('9. Funzionalità avanzate')

H2('9.1 Slot localStorage A/B/C/D')
P(
'Quattro slot di salvataggio rapido nella sezione STORAGE. Tap breve su un bottone carica '
'lo slot, tieni premuto per mezzo secondo per salvare lo stato corrente. Lo slot mostra un '
'pallino di conferma quando è popolato.'
)
P(
'Il contenuto salvato include: pattern drum, pattern bass step, live loop bass, parametri '
'di traccia drum, parametri di traccia bass, BPM, swing, lunghezza pattern, song sequence, '
'master mixer, modalità bass (step/looper), Performance Mode.'
)
P(
'I dati sono persistiti in localStorage, quindi sopravvivono a reload e chiusure dell\'app, '
'ma non a un cambio di browser o a un wipe del profilo.'
)

H2('9.2 Export / Import JSON v3')
P(
'Il bottone EXPORT scarica il set corrente come file .json in formato v3 (vedi appendice). '
'IMPORT apre un file picker per caricare un .json esportato in precedenza, una demo della '
'libreria, o un set ricevuto da altri.'
)
P(
'I file v1 e v2 (DrumAPP originale) si caricano normalmente; il blocco bass viene '
'inizializzato vuoto. I file v3 senza il campo performanceMode (le cinque demo bass '
'storiche) caricano in CLASSIC.'
)

H2('9.3 SHARE link')
P('Il bottone SHARE genera un URL compatto e lo copia negli appunti. L\'URL include:')
bullet('il pattern drum corrente (encoding hex bitmap, sedici step massimo);')
bullet('BPM, master drum, master bass;')
bullet('il pattern bass step corrente (base64 url-safe del JSON dell\'array di step).')
P(
'Lo SHARE link non include il live loop, i parametri di traccia, gli altri pattern, la '
'song sequence, la Performance Mode. È pensato come "biglietto da visita" del groove '
'corrente, non come backup completo. Per backup usa EXPORT.'
)

H2('9.4 BOUNCE WAV')
P('Il bottone BOUNCE apre un modale con quattro opzioni:')
bullet('2 loop del pattern corrente')
bullet('4 loop del pattern corrente')
bullet('8 loop del pattern corrente')
bullet('intera song (tutta la sequence A/A/B/A…)')
P(
'Il render usa OfflineAudioContext con sample rate 44.100 Hz, 2 canali, PCM 16-bit, output '
'.wav RIFF. La catena audio offline replica fedelmente quella runtime: stesse voci drum, '
'stesso supersaw stack del basso, stessa modalità performance attiva, stesso bus '
'drum/bass con HPF 30 Hz, stesso master gain. Il file scaricato è quindi identico a quello '
'che ascolti in tempo reale.'
)
P('Il bounce è deterministico: lo stesso set produce sempre lo stesso file.')

H2('9.5 REC live')
P(
'Il bottone REC registra tutto ciò che senti, inclusi knob twist, pattern switch e '
'movimenti dei master. Tecnicamente usa MediaStreamDestination in parallelo a '
'destination; il MediaRecorder autonegozia il codec (di default webm/opus, fallback mp4 '
'su Safari).'
)
P('Differenza chiave rispetto al BOUNCE:')
bullet('BOUNCE è offline e deterministico; cattura solo il pattern (o la song) statica.')
bullet('REC live è in tempo reale; cattura le tue interazioni, è ideale per performance live o sketches con automazioni.')

H2('9.6 Web MIDI out')
P(
'Il bottone MIDI apre requestMIDIAccess e si aggancia alla prima porta MIDI out '
'disponibile. Da quel momento ogni step drum suonato invia note General MIDI sul canale 10:'
)
table(
    ['Voce', 'MIDI note'],
    [
        ('KICK',    '36'),
        ('SNARE',   '38'),
        ('HI-HAT',  '42'),
        ('OPEN HH', '46'),
        ('CLAP',    '39'),
        ('TOM',     '45'),
        ('RIMSHOT', '37'),
        ('COWBELL', '56'),
    ],
    col_widths=[5, 4]
)
P('Le note bass vanno sul canale 2 con la nota suonata (C1 = 24, C2 = 36, eccetera).')
P('L\'integrazione MIDI è "out only": l\'app non riceve MIDI, ma può triggerare strumenti hardware o software esterni.')

H2('9.7 Undo / Redo')
Pkbd(
'Storico di quaranta stati. `⌘Z` (Mac) o `Ctrl+Z` annulla; `⌘⇧Z` o `Ctrl+Shift+Z` rifà. '
'Lo storico cattura cambi di pattern, parametri di traccia, BPM, swing, modalità bass, '
'Performance Mode, master mixer e song sequence.'
)

H2('9.8 Copy / Paste pattern')
P(
'Il bottone COPY copia il pattern drum corrente in un buffer interno; PASTE lo incolla sul '
'pattern corrente (sovrascrivendolo). Utile per duplicare A in B e modificarlo.'
)
P('Nota: COPY/PASTE agisce sul pattern drum, non sul basso né sui parametri di traccia.')

H2('9.9 Tap tempo')
Pkbd(
'`T` o il bottone TAP nel pannello tempo: ogni tocco registra un timestamp; dopo tre tocchi '
'l\'app calcola il BPM medio e lo applica in tempo reale. Usalo per allineare il tempo a un '
'brano che stai ascoltando.'
)

H2('9.10 Humanize e metronomo')
P(
'HUMANIZE (toggle nella modebar) applica micro-variazioni casuali di timing e velocity '
'ad ogni step, su drum e bass insieme. METRO attiva un click in alta frequenza sul primo '
'beat e in bassa frequenza sugli altri, utile per pratica e registrazione.'
)

page_break()

# ============================================================
# 10. SHORTCUT
# ============================================================
H1('10. Shortcut tastiera')

H2('10.1 Globali')
table(
    ['Tasto', 'Azione'],
    [
        ('SPACE',                'Play / Stop'),
        ('T',                    'Tap tempo'),
        ('1 – 4',                'Switch pattern A/B/C/D'),
        ('[ / ]',                'Pattern precedente / successivo'),
        ('↑ / ↓',                'Traccia attiva precedente / successiva'),
        ('M / S',                'Mute / Solo sulla traccia attiva'),
        ('C',                    'Clear pattern corrente (drum + bass)'),
        ('D',                    'Load demo'),
        ('B',                    'Toggle focus drum ⇄ bass nel pannello editing'),
        ('L',                    'Toggle modalità bass STEP ⇄ LOOPER'),
        ('P',                    'Cycla Performance Mode CLASSIC → WIDE → PUNCH → SUB'),
        ('⌘Z / Ctrl+Z',          'Undo'),
        ('⌘⇧Z / Ctrl+Shift+Z',   'Redo'),
        ('?',                    'Apri guida integrata'),
        ('ESC',                  'Chiudi modali'),
    ],
    col_widths=[5, 11.5]
)

H2('10.2 Solo in modalità LOOPER bass')
P('Queste shortcut non sono attive in STEP mode, per non collidere con quelle globali. Si attivano automaticamente quando il bass passa in LOOPER (L).')
table(
    ['Tasto', 'Azione'],
    [
        ('R',                          'Arma / disarma REC del loop bass'),
        ('A W S E D F T G Y H U J K', 'Mappatura pianistica delle dodici note (più la C alta su K)'),
        ('Z / X',                      'Ottava giù / su della tastiera on-screen'),
    ],
    col_widths=[5.5, 11]
)
P(
'Le note pianistiche partono dalla C dell\'ottava corrente (default 2). A = C, W = C#, '
'S = D, E = D#, D = E, F = F, T = F#, G = G, Y = G#, H = A, U = A#, J = B, K = C ottava sopra.'
)

page_break()

# ============================================================
# 11. ROADMAP
# ============================================================
H1('11. Feature roadmap')

quote(
'Le idee elencate in questa sezione sono in valutazione, non promesse. Non hanno date di '
'rilascio né ordini di priorità dichiarati. Sono dichiarate qui per dare alla community '
'early-adopter una finestra sulla direzione del progetto e per raccogliere feedback su '
'cosa risuona di più con l\'uso reale.'
)

H2('11.1 DrumAPPBass-Studio: sample-based bass')
P(
'Un fork separato (non sostitutivo dell\'attuale) con bassline costruite a partire da '
'sample WAV multi-velocity layer. Obiettivo: superare il soffitto onesto della sintesi '
'vanilla (≈ 7,5 / 10) e arrivare a un suono DAW-like (≈ 9 / 10), pur restando in ambito '
'web. Trade-off: il pacchetto cresce di un ordine di grandezza (decine di MB di sample) e '
'perde la portabilità "load offline in due secondi".'
)
P(
'L\'idea è di tenere DrumAPPBass come "sketch tool" leggero e DrumAPPBass-Studio come '
'"tool da finalizzazione", entrambi parte della stessa serie Weekend App.'
)

H2('11.2 Filtro per voce in WIDE mode')
P(
'In WIDE mode lo spread stereo è oggi sulle voci sawtooth ma il BiquadFilter è ancora '
'globale (un singolo nodo per tutte le voci). Aggiungere cinque filtri biquad indipendenti, '
'uno per voce, con leggere variazioni di cutoff e Q, darebbe spazialità anche al filtro '
'stesso e renderebbe lo spread più "tridimensionale". Costo CPU stimato: +30–50 % sulla '
'sezione bass.'
)

H2('11.3 Modalità FM / PWM')
P(
'Le quattro modalità attuali variano detune, panning, mix sub, drive ed envelope, ma la '
'forma d\'onda primaria resta sawtooth in tutte. Una quinta modalità con waveform diversa '
'(per esempio square con PWM, oppure FM a due operatori) introdurrebbe un carattere '
'radicalmente nuovo. Esempio: una mode "BUZZ" con onda quadra modulata in pulse-width per '
'un basso più 8-bit / chiptune.'
)

H2('11.4 Compressore master sul bass bus')
P(
'Aggiungere un DynamicsCompressorNode sul bassBus (con preset glue: ratio 3:1, threshold '
'−18 dB, attack 5 ms, release 100 ms) per dare quel "collante" tipico dei mix '
'professionali, soprattutto in PUNCH e SUB. È una modifica leggera e compatibile con il '
'bounce offline.'
)

H2('11.5 Espansione delle demo per modalità')
P(
'Oggi ogni Performance Mode ha una sola demo rappresentativa. Espandere a tre demo per '
'modalità (per esempio: WIDE in three flavours — synth-pop, electro funk, synthwave veloce '
'— con tonalità e BPM diversi) renderebbe più immediato l\'esplorare il "what if" del '
'singolo scenario.'
)

H2('11.6 MIDI in (input esterno)')
P(
'Oggi l\'integrazione MIDI è solo out. Aggiungere MIDI in permetterebbe di pilotare il '
'basso da una tastiera hardware, e il drum da pad esterni. Il sequencer interno resterebbe '
'come modalità predefinita, ma l\'utente potrebbe registrare il loop bass (e in futuro '
'anche il drum) suonandolo da hardware.'
)

H2('11.7 Mobile UX improvements')
P(
'Su iPhone e iPad alcune interazioni sono migliorabili: il long-press per aprire il '
'mini-selettore di nota va affinato (oggi confligge in alcuni casi con il drag verticale), '
'la tastiera on-screen del basso può guadagnare in altezza in landscape, il modale demos '
'può scrollare con più fluidità su iOS Safari. È un\'area di raffinamento continuo che '
'dipende dal feedback degli early adopter.'
)

H2('11.8 Versione internazionale del manuale')
P(
'Questo documento esiste oggi solo in italiano. Versioni inglese, spagnola e francese '
'sono in valutazione per la fase di rilascio pubblico, in modo da estendere l\'adozione '
'fuori dalla community italiana iniziale.'
)

H2('11.9 Esportazioni alternative')
P(
'Oltre al WAV, sono in valutazione altri formati di export: stem WAV separati (drum-only '
'e bass-only in due file), MIDI file dei pattern (.mid invece del solo MIDI live), '
'Ableton Live .als semplificato. Ognuna di queste è una piccola feature ma ad alto valore '
'per chi usa DrumAPPBass come prefase di una produzione su DAW.'
)

H2('11.10 Engagement community')
P(
'L\'autore sta valutando — sulla base del feedback ricevuto durante la fase di anteprima — '
'di aprire un canale Telegram o un forum Discord dedicato per chi usa l\'app, dove '
'condividere demo, fare contest periodici (per esempio "best 30-second bassline del mese"), '
'e raccogliere bug report e idee in modo più strutturato rispetto a WhatsApp.'
)

page_break()

# ============================================================
# 12. FAQ
# ============================================================
H1('12. FAQ')

faqs = [
('L\'app si scarica oppure si usa solo online?',
 'Si usa nel browser, ma essendo una PWA viene cachata al primo accesso (via service worker '
 'drumappbass-v2). Da quel momento funziona offline finché il browser non svuota la cache. '
 'Su mobile si può "installare" alla home screen e si comporta come un\'app standalone.'),
('Devo creare un account?',
 'No. Nessun account, nessun login, nessun tracking, nessuna telemetria. I dati sono '
 'salvati solo in localStorage sul tuo dispositivo.'),
('Funziona su iPhone? Su Android?',
 'Sì, su entrambi. Su iOS Safari il MediaRecorder usa mp4 (fallback automatico), su Chrome '
 'Android usa webm/opus. Su iOS è consigliato installare la PWA alla home per poter '
 'sbloccare il contesto audio in modo affidabile.'),
('Posso usare DrumAPPBass dal vivo?',
 'Sì, è uno degli use case principali. Il REC live cattura la performance, lo SHARE link è '
 'utile per condividere groove al volo. Per palco serio si consiglia di usare l\'app '
 'installata come PWA, in modalità aereo se non serve internet, e con cuffie cablate (il '
 'Bluetooth introduce latenza che non è gestibile lato browser).'),
('Posso usare le demo per pubblicare un mio brano?',
 'Sì. Le demo sono ricostruzioni ritmico-armoniche generiche, non sample di brani '
 'originali, e il codice e gli output dell\'app sono coperti da licenza MIT. Il tuo brano '
 'finale (esportato come WAV o sviluppato in DAW partendo dal MIDI / WAV bounce) è opera '
 'tua a tutti gli effetti.'),
('Funziona offline?',
 'Sì, completamente. Il service worker pre-cacha tutti gli asset core al primo accesso. '
 'Le demo si caricano via fetch dalla cartella examples/: se la cache contiene già i file, '
 'funzionano anche offline.'),
('Come aggiorno l\'app dopo un nuovo Round?',
 'Su desktop, hard reload con Cmd+Shift+R (Mac) o Ctrl+Shift+R (Win/Linux). Su iOS, '
 'chiudi e riapri la PWA due volte: la prima volta scarica la nuova versione, la seconda '
 'l\'attiva.'),
('Ho trovato un bug. Come segnalo?',
 'Vedi sezione 13 (Risorse). I canali sono GitHub Issues, WhatsApp diretto all\'autore '
 '(per chi è nella community early-adopter), email.'),
('Posso modificare il codice?',
 'Sì, la licenza MIT lo permette. Fork del repository, modifica in locale, eventualmente '
 'apri una pull request se vuoi contribuire al progetto principale.'),
('C\'è una versione Mac/Windows nativa?',
 'No, e non è prevista. La filosofia "vive nel browser" è un pilastro del progetto: una '
 'PWA si installa come app ma resta zero-dipendenze e cross-platform.'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    r = p.add_run(q)
    r.bold = True; r.font.color.rgb = BASS
    P(a)

page_break()

# ============================================================
# 13. RISORSE
# ============================================================
H1('13. Risorse, licenza e contatti')

H2('13.1 Link')
bullet('App live: pezzaliapp.github.io/DrumAPPBass — versione canonica, sempre aggiornata.')
bullet('Repository: github.com/pezzaliapp/DrumAPPBass — codice sorgente, demo, changelog.')
bullet('Sito autore: pezzaliapp.com')
bullet('DrumAPP originale: github.com/pezzaliapp/DrumAPP — il progetto da cui DrumAPPBass è stato forkato.')

H2('13.2 Come dare feedback')
P('In ordine di preferenza, a seconda del contesto.')
bullet('Bug riproducibili o richieste tecniche: apri una issue su GitHub (Issues → New issue). Includi browser, sistema operativo, passi per riprodurre.')
bullet('Feedback informale, idee, prime impressioni: WhatsApp diretto all\'autore se sei nella community early-adopter; email per chi non è in community.')
bullet('Pull request: benvenute, soprattutto se fixano un bug riproducibile. Per feature nuove di una certa portata, consiglio di aprire prima una issue e discutere l\'approccio.')

H2('13.3 Licenza')
P('DrumAPPBass è distribuito sotto licenza MIT.')
P(
'In sintesi: puoi usare, copiare, modificare, ridistribuire, vendere il software o opere '
'derivate, a patto di mantenere il copyright notice e la licenza nei file. Il software è '
'fornito "as is", senza garanzie. Vedi il file LICENSE nel repository per il testo integrale.'
)

H2('13.4 Note legali')

H4('Demo e copyright musicale')
P(
'Le demo della libreria sono ricostruzioni ritmico-armoniche generiche ispirate ai loro '
'generi di riferimento. Non sono trascrizioni né copie di brani specifici, non riproducono '
'melodie identificative, non utilizzano campioni audio originali. I nomi tipo "Billie '
'Jean-style" o "Seven Nation Army" indicano lo stile, non il brano. Ogni file generato '
'dall\'utente con DrumAPPBass è opera dell\'utente.'
)

H4('Tracking e privacy')
P(
'L\'applicazione non raccoglie alcun dato dell\'utente. Non c\'è analytics, non c\'è '
'telemetria, non c\'è tracker. Tutti i dati di salvataggio (slot, set correnti) sono in '
'localStorage locale e non lasciano mai il dispositivo. Lo SHARE link contiene solo il '
'pattern serializzato in URL hash; non passa per server di terze parti.'
)

H4('Vanilla Web Audio')
P(
'L\'audio è generato in tempo reale via Web Audio API standard. Non ci sono librerie audio '
'esterne (Tone.js, Howler, eccetera). Non ci sono sample WAV inclusi o caricati. Tutto il '
'suono è risultato di sintesi diretta dei nodi Oscillator, Gain, BiquadFilter, WaveShaper, '
'BufferSource con noise pre-generato.'
)

page_break()

# ============================================================
# 14. APPENDICE
# ============================================================
H1('14. Appendice tecnica')

H2('14.1 Architettura della sintesi')
P('Tutta la sintesi vive in app.js. Le sezioni principali sono commentate con header tipo // 5b) SINTESI BASS.')

H4('Catena drum (per ciascuna delle otto tracce)')
code('voice -> trackPanner[i] -> trackFilter[i] -> trackGain[i] -> drumBus -> masterGain -> destination')

H4('Catena bass')
code(
'note -> [stack supersaw 5 voci sawtooth detunate] -> stackMix\n'
'note -> [sub-oscillator square -12 st (condizionale)]\n'
'                                    |\n'
'                            mixSaw + mixSub\n'
'                                    |\n'
'                        gate (envelope ampiezza)\n'
'                                    |\n'
'                          bassPanner (globale)\n'
'                                    |\n'
'                  bassFilter (LP, env+Q dinamici per nota)\n'
'                                    |\n'
'                     bassDrive (WaveShaper tanh)\n'
'                                    |\n'
'                     bassGain (volume traccia)\n'
'                                    |\n'
'                 bassBus -> bassBusHP (30 Hz) -> masterGain'
)

H4('Scheduler')
P(
'Pattern di Chris Wilson: un setInterval ogni 25 ms programma gli step nei prossimi 100 ms '
'via AudioContext.currentTime. Il bass step sequencer si innesta nello stesso scheduler, '
'nello stesso scheduleStep(stepIdx, baseTime). Il bass looper schedula le note dal proprio '
'array bassLiveLoop[pattern] in base alla posizione frazionaria dentro il pattern.'
)

H4('Bounce offline')
P(
'OfflineAudioContext ricostruisce la stessa identica catena audio (incluso il drive boost '
'e il resonance boost della Performance Mode attiva) e schedula tutti gli step della '
'sequenza scelta. Output PCM 16-bit stereo encodato come WAV RIFF.'
)

H2('14.2 Formato JSON v3')
code('''{
  "version": 3,
  "exportedAt": "2026-04-25T12:00:00Z",
  "bpm": 110,
  "swing": 0,
  "patternLength": 16,
  "humanize": false,
  "masterDrum": 0.9,
  "masterBass": 0.8,
  "performanceMode": "wide",
  "trackParams": [ /* 8 tracce drum */ ],
  "patterns": {
    "A": [ /* 8 array da 16 celle */ ],
    "B": [...], "C": [...], "D": [...]
  },
  "songSequence": ["A", "A", "B", "A"],
  "bass": {
    "trackParams": {
      "volume": 0.75, "pan": 0, "cutoff": 0.55,
      "resonance": 3.0, "envAmount": 0.55,
      "decay": 160, "drive": 0.18,
      "mute": false, "solo": false
    },
    "patterns": {
      "A": [ {"note":"B1","vel":0.9,"len":0.30,
              "accent":true,"slide":false},
             null, /* ... 14 celle ... */ ],
      "B": [...], "C": [...], "D": [...]
    },
    "liveLoops": {
      "A": [ {"step":0.0,"note":"B1","vel":0.9,"len":0.30} ],
      "B": [], "C": [], "D": []
    },
    "mode": "step"
  }
}''')

P('Backward compatibility:')
bullet('File version < 3 (DrumAPP pre-fork) → si carica normalmente, sezione bass inizializzata vuota.')
bullet('File version 3 senza performanceMode → carica in classic (caso delle cinque demo bass storiche e di qualunque export pre-Round 7).')
bullet('File con performanceMode non noto → fallback classic con ignorare silenzioso.')
P(
'Il numero di versione resta 3 anche dopo l\'aggiunta di performanceMode: il campo è '
'additivo e non rompe il parser dei consumer esistenti.'
)

H2('14.3 Performance Modes — dettaglio interno')
P('Ciascuna modalità è un oggetto PERFORMANCE_MODES[key] con i seguenti campi:')
table(
    ['Campo', 'Tipo', 'Significato'],
    [
        ('label',            'string',   'Nome visualizzato (CLASSIC, WIDE, …)'),
        ('desc',             'string',   'Descrizione one-line per UI'),
        ('detunes',          'number[]', 'Array di cents per voce (5 o 3 elementi)'),
        ('pans',             'number[]', 'Array di pan –1.0..+1.0 (stessa lunghezza di detunes)'),
        ('gains',            'number[]', 'Array di gain normalizzati (somma ≈ 1)'),
        ('subEnabled',       'bool',     'Se attivare il sub-oscillatore'),
        ('subMixRatio',      '0..1',     'Gain del sub nel mixer'),
        ('sawMixRatio',      '0..1',     'Gain dello stack saw nel mixer'),
        ('subMidiThreshold', 'int',      'MIDI minima sotto cui il sub si disabilita (36 = C2, 35 = B1)'),
        ('driveBoost',       'number',   'Moltiplicatore di bassParams.drive'),
        ('cutoffOffset',     '−1..+1',   'Offset additivo su bassParams.cutoff'),
        ('resonanceBoost',   'number',   'Moltiplicatore di bassParams.resonance'),
        ('envAmountBoost',   'number',   'Moltiplicatore di bassParams.envAmount'),
        ('sustainLevel',     '0..1',     'Sustain del gate envelope amp'),
        ('attackSec',        'sec',      'Attack del gate envelope amp'),
        ('releaseSec',       'sec',      'Release del gate envelope amp'),
    ],
    col_widths=[4, 2.5, 10]
)
P('I valori esatti delle quattro modalità sono nel file app.js, sezione PERFORMANCE_MODES.')

H2('14.4 Limiti noti')
bullet('Un solo basso monofonico (per scelta di design, non limitazione tecnica).')
bullet('Pattern multipli A/B/C/D, non "scenes" come in Live (16 scene). Per song lunghe si compone in song sequence.')
bullet('Lo SHARE link include solo il pattern corrente, non gli altri tre. Per backup completo usare EXPORT JSON.')
bullet('Web MIDI è "out only", non c\'è MIDI in.')
bullet('Il REC live su iOS Safari produce mp4 invece di webm/opus (limitazione del browser).')
bullet('Il bounce WAV su pattern molto lunghi (intera song con 16 slot e patternLength 32) può richiedere 10-20 secondi di rendering: non c\'è una progress bar dettagliata, solo un toast "Bounce in corso…".')

H2('14.5 Cambio di versione')
table(
    ['Versione', 'Data', 'Highlight'],
    [
        ('v1.1', '2026-04-25', 'Round 7 — Performance Modes (CLASSIC / WIDE / PUNCH / SUB), 4 nuove demo, shortcut P'),
        ('v1.0', '2026-04-24', 'Prima release pubblica DrumAPPBass — fork DrumAPP con sezione bass completa'),
    ],
    col_widths=[2.5, 3, 11]
)

# Footer
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Fine del manuale. Buon uso e buoni groove.')
r.italic = True; r.font.color.rgb = GREY; r.font.size = Pt(11)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— Alessandro Pezzali')
r.italic = True; r.font.color.rgb = GREY; r.font.size = Pt(11)

# Save
doc.save(OUT)
print(f'wrote {OUT}')
print(f'size: {os.path.getsize(OUT)} bytes')
